#!/usr/bin/env python3
"""Generate the cover-letter .docx copy-paste source from a markdown letter file.

Formatting per the formatting spec (PRIVATE__YOUR_FILES_GITIGNORED/04-TAILOR__YOUR_PRIVATE_INFO/cover-letter/formatting-spec.md): body font,
size, and text color come from config.json next to this script (gitignored instance; copy
config.template.json). Single spacing with space-after, bullets with bold lead-ins, real
inline hyperlinks (underlined, body-colored). The defaults below apply when no config exists.

Markdown subset understood:
  - blank-line-separated paragraphs
  - "- " bullets
  - **bold** runs
  - [anchor](url) links (may contain **bold** inside anchor — not needed, kept simple)

Usage (from the repo root):
  .venv/bin/python3 ENGINE__PUBLIC_GIT_TRACKED/04-TAILOR/cover-letter/make_cover_letter_docx.py letter.md -o out.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is not installed. Run: .venv/bin/pip install python-docx", file=sys.stderr)
    sys.exit(2)

# Defaults; overridden by the "docx" section of config.json when present.
DEFAULTS = {
    "font": "Helvetica Neue", "size_pt": 10, "color_hex": "3F3F3F",
    # Bullet blocks are INDENTED as a group, not flush with the body prose: the bullet
    # glyph sits at `bullet_indent_in` and the text at `bullet_text_indent_in`, so
    # wrapped lines align under the text rather than under the "•". Without the group
    # indent the list reads as more body paragraphs; with it, the letter has a visible
    # scannable middle. The gap between the glyph and the text is the DIFFERENCE of the
    # two (0.375 - 0.25 = 0.125"); a wider gap reads as a gutter, so keep the text indent
    # close to the glyph. `bullet_space_after_pt` is the gap BETWEEN bullets.
    "bullet_indent_in": 0.25,
    "bullet_text_indent_in": 0.375,
    "bullet_space_after_pt": 8,
}


def load_docx_config(path):
    cfg = dict(DEFAULTS)
    if path and Path(path).is_file():
        try:
            user = json.loads(Path(path).read_text(encoding="utf-8"))
        except (OSError, ValueError) as e:
            print(f"cannot read config {path}: {e}", file=sys.stderr)
            sys.exit(2)
        cfg.update(user.get("docx", {}))
    cfg["color_hex"] = cfg["color_hex"].lstrip("#").upper()
    return cfg


def _set_tab_stop(paragraph, position):
    """Pin an explicit left tab stop at `position` for the bullet's "•\\t".

    Without it Word falls back to its default half-inch grid, so the text lands
    wherever that grid happens to sit rather than at the hanging indent — the bullet
    and its wrapped lines then disagree, which is exactly the ragged look this
    formatting is meant to remove.
    """
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(position.twips)))
    tabs.append(tab)
    paragraph._p.get_or_add_pPr().append(tabs)


INLINE = re.compile(r"(\[([^\]]+)\]\(([^)\s]+)\))|(\*\*([^*]+)\*\*)")


def smarten(text):
    """Typographic quotes/apostrophes/ellipses so the paste into the editor is clean."""
    text = re.sub(r"(\w)'(\w)", "’".join((r"\1", r"\2")), text)   # don't -> don’t
    text = re.sub(r"'(\w)", "‘" + r"\1", text)                     # opening single
    text = re.sub(r"(\w)'", r"\1" + "’", text)                     # trailing
    text = re.sub(r'"(\S[^"]*?)"', "“" + r"\1" + "”", text)   # double pairs
    return text


class DocBuilder:
    def __init__(self, cfg):
        self.font = cfg["font"]
        self.size = Pt(cfg["size_pt"])
        self.color = RGBColor.from_string(cfg["color_hex"])
        self.color_hex = cfg["color_hex"]
        self.bullet_indent = Inches(float(cfg["bullet_indent_in"]))
        self.bullet_text_indent = Inches(float(cfg["bullet_text_indent_in"]))
        self.bullet_space_after = Pt(float(cfg["bullet_space_after_pt"]))

    def style_run(self, run, bold=False, underline=False):
        run.font.name = self.font
        run.font.size = self.size
        run.font.color.rgb = self.color
        run.bold = bold
        run.underline = underline
        # ensure the East Asian / complex-script font names match too (Word quirk)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rFonts.set(qn(attr), self.font)

    def add_hyperlink(self, paragraph, text, url):
        """Real hyperlink relationship + underlined run in the body color."""
        part = paragraph.part
        r_id = part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        color = OxmlElement("w:color"); color.set(qn("w:val"), self.color_hex); rPr.append(color)
        new_run.append(rPr)
        t = OxmlElement("w:t"); t.text = text
        t.set(qn("xml:space"), "preserve")
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        # style the run through the python-docx wrapper for font/size
        from docx.text.run import Run
        run = Run(new_run, paragraph)
        self.style_run(run, underline=True)
        return run

    def add_inline(self, paragraph, text):
        pos = 0
        for m in INLINE.finditer(text):
            if m.start() > pos:
                self.style_run(paragraph.add_run(smarten(text[pos:m.start()])))
            if m.group(1):  # link
                self.add_hyperlink(paragraph, smarten(m.group(2)), m.group(3))
            else:           # bold
                self.style_run(paragraph.add_run(smarten(m.group(5))), bold=True)
            pos = m.end()
        if pos < len(text):
            self.style_run(paragraph.add_run(smarten(text[pos:])))

    def build(self, md_path, out_path):
        raw = Path(md_path).read_text(encoding="utf-8")
        # letters may carry an annotation header above "---"; use only the body below it
        if "\n---\n" in raw:
            head, body = raw.split("\n---\n", 1)
            # only strip if the head looks like an annotation (starts with '#')
            raw = body if head.lstrip().startswith("#") else raw

        doc = docx.Document()
        style = doc.styles["Normal"]
        style.font.name = self.font
        style.font.size = self.size
        style.font.color.rgb = self.color
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(10)
        style.paragraph_format.line_spacing = 1.0

        blocks = [b.strip("\n") for b in re.split(r"\n\s*\n", raw) if b.strip()]
        for block in blocks:
            lines = block.splitlines()
            if all(ln.lstrip().startswith(("- ", "• ")) for ln in lines):
                for ln in lines:
                    p = doc.add_paragraph()
                    # Text at the deeper indent; the bullet hangs back to the shallower
                    # one (a NEGATIVE first-line indent of exactly the difference), so
                    # wrapped lines align under the text, not under the glyph.
                    p.paragraph_format.left_indent = self.bullet_text_indent
                    p.paragraph_format.first_line_indent = -(
                        self.bullet_text_indent - self.bullet_indent)
                    p.paragraph_format.space_after = self.bullet_space_after
                    p.paragraph_format.line_spacing = 1.0
                    _set_tab_stop(p, self.bullet_text_indent)
                    self.style_run(p.add_run("•\t"))
                    self.add_inline(p, ln.lstrip()[2:].strip())
            else:
                # signature-style short adjacent lines (e.g. "Looking forward to chatting,\n<Name>")
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                for i, ln in enumerate(lines):
                    if i:
                        p.add_run().add_break()
                    self.add_inline(p, ln)

        doc.save(out_path)
        return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("md_file")
    ap.add_argument("-o", "--out", required=True)
    ap.add_argument("--config", default=str(Path(__file__).resolve().parents[3] / "PRIVATE__YOUR_FILES_GITIGNORED" / "04-TAILOR__YOUR_PRIVATE_INFO" / "cover-letter" / "config.json"),
                    help="config JSON (default: config.json beside this script; falls back to defaults)")
    args = ap.parse_args()
    cfg = load_docx_config(args.config)
    out = DocBuilder(cfg).build(args.md_file, args.out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
