"""Cover-letter .docx bullet formatting.

Bullet blocks are indented as a GROUP (glyph at one stop, text at a deeper one) so a
letter has a visible, scannable middle instead of reading as more body paragraphs. This
was hand-fixed in Word on every letter before it was automated; these tests keep it
automated.
"""
import importlib.util
import sys
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")

REPO = Path(__file__).resolve().parents[1]
MODULE = REPO / "ENGINE__PUBLIC_GIT_TRACKED" / "04-TAILOR" / "cover-letter" / "make_cover_letter_docx.py"

_spec = importlib.util.spec_from_file_location("make_cover_letter_docx", MODULE)
mcl = importlib.util.module_from_spec(_spec)
sys.modules["make_cover_letter_docx"] = mcl
_spec.loader.exec_module(mcl)

LETTER = """August 6, 2026

Hi there,

Here's what I'd bring:

- **First point.** A sentence long enough that it must wrap onto a second line so the hanging indent is actually exercised rather than merely configured.
- **Second point.** Another one, also with an [inline link](https://example.com) in it.

That's the pitch.

Looking forward to chatting,
Jessica
"""


def _build(tmp_path, cfg_overrides=None):
    md = tmp_path / "final.md"
    md.write_text(LETTER, encoding="utf-8")
    cfg = dict(mcl.DEFAULTS)
    cfg.update(cfg_overrides or {})
    out = tmp_path / "letter.docx"
    mcl.DocBuilder(cfg).build(str(md), str(out))
    return docx.Document(str(out))


def _bullets(doc):
    return [p for p in doc.paragraphs if p.text.lstrip().startswith("•")]


def _body(doc):
    return [p for p in doc.paragraphs if not p.text.lstrip().startswith("•")]


def test_bullets_are_indented_as_a_group_not_flush_with_the_body(tmp_path):
    doc = _build(tmp_path)
    bullets = _bullets(doc)
    assert len(bullets) == 2
    for p in bullets:
        # Text sits at the deeper stop; the glyph hangs back to the shallower one.
        assert p.paragraph_format.left_indent.inches == pytest.approx(0.5)
        assert p.paragraph_format.first_line_indent.inches == pytest.approx(-0.25)
    # Body prose stays flush left — only the list is indented.
    for p in _body(doc):
        assert not p.paragraph_format.left_indent


def test_hanging_indent_equals_the_gap_so_wrapped_lines_align_under_the_text(tmp_path):
    """If these two disagree, wrapped lines land under the bullet glyph instead of
    under the text — the exact ragged look the indenting is meant to remove."""
    doc = _build(tmp_path)
    for p in _bullets(doc):
        gap = p.paragraph_format.left_indent.inches - 0.25
        assert -p.paragraph_format.first_line_indent.inches == pytest.approx(gap)


def test_each_bullet_pins_an_explicit_tab_stop_at_the_text_indent(tmp_path):
    """Without an explicit stop Word falls back to its default half-inch grid and the
    text lands wherever that grid sits, not at the hanging indent."""
    from docx.oxml.ns import qn
    doc = _build(tmp_path)
    for p in _bullets(doc):
        stops = p._p.find(qn("w:pPr")).findall(qn("w:tabs") + "/" + qn("w:tab"))
        assert stops, "bullet paragraph has no explicit tab stop"
        # 0.5 inch == 720 twips
        assert [s.get(qn("w:pos")) for s in stops] == ["720"]


def test_bullets_are_separated_from_each_other(tmp_path):
    doc = _build(tmp_path)
    for p in _bullets(doc):
        assert p.paragraph_format.space_after.pt == pytest.approx(8)


def test_bullet_geometry_is_configurable(tmp_path):
    doc = _build(tmp_path, {"bullet_indent_in": 0.4, "bullet_text_indent_in": 0.9,
                            "bullet_space_after_pt": 4})
    for p in _bullets(doc):
        assert p.paragraph_format.left_indent.inches == pytest.approx(0.9)
        assert p.paragraph_format.first_line_indent.inches == pytest.approx(-0.5)
        assert p.paragraph_format.space_after.pt == pytest.approx(4)


def test_inline_links_and_bold_survive_the_indenting(tmp_path):
    """Formatting the paragraph must not flatten the runs inside it."""
    doc = _build(tmp_path)
    text = "\n".join(p.text for p in _bullets(doc))
    assert "First point." in text and "inline link" in text
    assert any(r.bold for p in _bullets(doc) for r in p.runs), "bold lead-in lost"
